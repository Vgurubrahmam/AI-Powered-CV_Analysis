"""DOCX text extractor — python-docx XML walker."""

from __future__ import annotations

import io
from dataclasses import dataclass

import structlog

log = structlog.get_logger(__name__)


@dataclass
class DOCXExtractResult:
    text: str
    char_count: int
    is_empty: bool
    method_used: str = "python-docx"
    table_count: int = 0


def extract_docx(file_bytes: bytes) -> DOCXExtractResult:
    """Extract text from DOCX bytes using python-docx."""
    try:
        from docx import Document
    except ImportError:
        log.warning("python_docx_not_installed")
        return DOCXExtractResult(text="", char_count=0, is_empty=True)

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        log.error("docx_open_failed", error=str(exc))
        return DOCXExtractResult(text="", char_count=0, is_empty=True)

    parts: list[str] = []
    table_count = 0

    # Walk paragraphs in document order
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Walk tables: left-to-right, top-to-bottom cell order
    for table in doc.tables:
        table_count += 1
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_cells.append(cell_text)
            if row_cells:
                parts.append(" | ".join(row_cells))

    full_text = "\n".join(parts)
    return DOCXExtractResult(
        text=full_text,
        char_count=len(full_text),
        is_empty=len(full_text.strip()) == 0,
        table_count=table_count,
    )
